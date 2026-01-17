import os
from PyPDF2 import PdfReader
from docx import Document


class FileParser:
    """Service for extracting text from uploaded files."""
    
    @staticmethod
    def extract_text(file_path: str) -> str:
        """
        Extract text content from a file based on its extension.
        
        Args:
            file_path: Path to the file
            
        Returns:
            Extracted text content
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        extension = os.path.splitext(file_path)[1].lower()
        
        if extension == '.pdf':
            return FileParser._extract_from_pdf(file_path)
        elif extension == '.docx':
            return FileParser._extract_from_docx(file_path)
        elif extension == '.txt':
            return FileParser._extract_from_txt(file_path)
        else:
            raise ValueError(f"Unsupported file type: {extension}")
    
    @staticmethod
    def _extract_from_pdf(file_path: str) -> str:
        """Extract text from a PDF file."""
        try:
            reader = PdfReader(file_path)
            text_parts = []
            
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
            
            return '\n'.join(text_parts)
        except Exception as e:
            raise Exception(f"Error reading PDF: {str(e)}")
    
    @staticmethod
    def _extract_from_docx(file_path: str) -> str:
        """Extract text from a DOCX file."""
        try:
            doc = Document(file_path)
            text_parts = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        text_parts.append(row_text)
            
            return '\n'.join(text_parts)
        except Exception as e:
            raise Exception(f"Error reading DOCX: {str(e)}")
    
    @staticmethod
    def _extract_from_txt(file_path: str) -> str:
        """Extract text from a TXT file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except UnicodeDecodeError:
            # Try with different encoding
            with open(file_path, 'r', encoding='latin-1') as f:
                return f.read()
        except Exception as e:
            raise Exception(f"Error reading TXT: {str(e)}")
    
    @staticmethod
    def allowed_file(filename: str, allowed_extensions: set) -> bool:
        """Check if a file has an allowed extension."""
        return '.' in filename and \
               filename.rsplit('.', 1)[1].lower() in allowed_extensions
